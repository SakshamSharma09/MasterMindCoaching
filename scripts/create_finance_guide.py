from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "finance-guide"
DOCX = OUT / "MasterMind Finance Operations Guide.docx"

NAVY = "0A1D39"
BLUE = "2E74B5"
TEAL = "159A83"
GOLD = "D9A12D"
PALE = "E8EEF5"
LIGHT = "F7FAFC"
MUTED = "5B677A"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color="C9D3E0", size="6"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)
        tbl_borders.append(node)


def fixed_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(table)


def set_font(run, size=11, color=NAVY, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def paragraph(doc, text="", *, bold=False, color=NAVY, align=None, after=6, line=1.25):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    size, color, before, after = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, "1F4D78", 10, 5),
    }[level]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=size, color=color, bold=True)
    return p


def callout(doc, title, body, fill="EAF7F4", accent=TEAL):
    table = doc.add_table(rows=1, cols=2)
    fixed_table(table, [0.18, 6.32])
    shade(table.cell(0, 0), accent)
    shade(table.cell(0, 1), fill)
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(title + "\n"), bold=True, color=NAVY)
    set_font(p.add_run(body), color=NAVY)
    return table


def step_table(doc, steps):
    table = doc.add_table(rows=len(steps), cols=2)
    fixed_table(table, [0.62, 5.88])
    for idx, (title, body) in enumerate(steps, 1):
        shade(table.cell(idx - 1, 0), NAVY)
        p0 = table.cell(idx - 1, 0).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p0.add_run(str(idx)), size=13, color=WHITE, bold=True)
        p = table.cell(idx - 1, 1).paragraphs[0]
        set_font(p.add_run(title + "\n"), bold=True)
        set_font(p.add_run(body), color=MUTED)
    return table


def grid(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    fixed_table(table, widths)
    for i, header in enumerate(headers):
        shade(table.cell(0, i), PALE)
        set_font(table.cell(0, i).paragraphs[0].add_run(header), bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_margins(cells[i])
            set_font(cells[i].paragraphs[0].add_run(str(value)), color=NAVY)
    return table


def page_break(doc):
    doc.add_page_break()


def page_number(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_end])


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
):
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(header.add_run("MASTERMIND COACHING CLASSES  •  FINANCE OPERATIONS"), size=9, color=MUTED, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run("Finance Operations Guide  •  "), size=9, color=MUTED)
page_number(footer)

# Editorial cover
paragraph(doc, "OPERATIONS GUIDE", bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
p = paragraph(doc, "MasterMind Finance\nOperations Guide", bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
for run in p.runs:
    run.font.size = Pt(28)
paragraph(doc, "Fees • Collections • Parent reminders • Salaries • Expenses • Reports", color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
callout(
    doc,
    "Release 1.0.12",
    "This guide describes the guided Finance workflow introduced with next-cycle billing and real PDF receipts.",
    fill="FFF8E8",
    accent=GOLD,
)
paragraph(doc, "\nPrepared for administrators of MasterMind Coaching Classes", color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
paragraph(doc, "28 July 2026", color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break(doc)
heading(doc, "1. Finance at a glance")
paragraph(doc, "Use one clear path for each job. Fees Management defines what a student owes. Fee Collection records money received. Expenses tracks salaries and all other costs.")
grid(doc, ["Need to do", "Go here", "Result"], [
    ("Assign or change a fee schedule", "Finance → Fees Management", "Installments generated by billing period"),
    ("Collect full or partial payment", "Finance → Fee Collection", "Payment record + PDF receipt"),
    ("Pay a teacher salary", "Finance → Expenses", "Salary paid + salary receipt"),
    ("Record rent, utilities, or supplies", "Finance → Expenses", "Pending/overdue/paid obligation"),
    ("Review totals", "Finance → Overview / Reports", "Revenue, expenses, pending balance, net profit"),
], [2.2, 1.9, 2.4])
heading(doc, "Quick start", 2)
step_table(doc, [
    ("Create or select a fee plan", "Choose Monthly, Quarterly, Half-yearly, Annual, or One-time."),
    ("Assign it to the student", "Enter billing start, first due date, and session end."),
    ("Collect a specific installment", "Choose the child and month/period, then record full or partial payment."),
    ("Send the receipt", "Download the PDF, share through WhatsApp, or email it when the parent has supplied an email."),
])

page_break(doc)
heading(doc, "2. Set up a student fee")
paragraph(doc, "All fee assignment happens in Fees Management. Fee Collection intentionally has no duplicate setup form.")
step_table(doc, [
    ("Open Add Fee", "Finance → Fees Management → Add Fee."),
    ("Select the student", "Only students from the selected academic session are listed."),
    ("Select a reusable fee plan", "The plan supplies frequency, default amount, and academic year."),
    ("Confirm the schedule", "Enter amount, billing start date, first due date, and session end date."),
    ("Save", "The current period is created once. Future periods appear automatically at their period start."),
])
heading(doc, "Frequency meanings", 2)
grid(doc, ["Frequency", "Cycle", "Example from 1 April"], [
    ("Monthly", "1 month", "April period due 1 May"),
    ("Quarterly", "3 months", "Apr–Jun period due 1 July"),
    ("Half-yearly", "6 months", "Apr–Sep period due 1 October"),
    ("Annual", "12 months", "Apr–Mar period due 1 April next year"),
    ("One-time", "No recurrence", "Due on the chosen date"),
], [1.35, 1.25, 3.9])
callout(doc, "Important", "Changing a plan later does not rewrite paid historical installments. Inactivating a student stops future periods but retains paid and partially paid history.", fill="FFF2F2", accent="D94C4C")

page_break(doc)
heading(doc, "3. Understand due and overdue dates")
paragraph(doc, "Recurring fees use next-cycle billing. A period is pending while it is being delivered and becomes overdue when its due date is reached.")
grid(doc, ["Date", "What exists", "Status"], [
    ("1 April", "April installment is created; due 1 May", "Pending"),
    ("30 April", "April installment still has a balance", "Pending"),
    ("1 May", "April reaches due date; May installment is created", "April Overdue; May Pending"),
    ("1 June", "May reaches due date; June installment is created", "May Overdue; June Pending"),
], [1.3, 3.7, 1.5])
heading(doc, "Schedule guardrails", 2)
grid(doc, ["Rule", "Behavior"], [
    ("Idempotency", "Opening Finance repeatedly cannot duplicate an installment."),
    ("Session end", "No period starts after the schedule end date."),
    ("Student inactivation", "The inactive date becomes the effective end for future generation."),
    ("Legacy rows", "Existing paid history and old due dates remain unchanged."),
    ("Outstanding balance", "Partial payment leaves the remaining amount visible and reminded."),
], [1.7, 4.8])

page_break(doc)
heading(doc, "4. Collect payment and issue a receipt")
step_table(doc, [
    ("Search the student", "Open Fee Collection and select the student from the real session list."),
    ("Select one or more installments", "Each row identifies the fee plan, period, due date, paid amount, and balance."),
    ("Enter the payment", "Choose full or partial amount, method, payment date, and transaction/reference."),
    ("Confirm collection", "The balance and status update in one transaction."),
    ("Distribute the receipt", "Download the PDF. WhatsApp opens the primary parent number; attach the PDF before sending. Email is available after the parent supplies an email."),
])
callout(doc, "Receipt controls", "A payment cannot exceed the installment balance. A PDF receipt is generated for every completed payment; quick “Mark Paid” shortcuts do not bypass this audit trail.")
heading(doc, "What the parent sees", 2)
paragraph(doc, "The Parent dashboard and Fees page show each pending or overdue installment by student, period, due date, and remaining balance. The reminder remains after a partial payment and disappears only after full payment, waiver, or cancellation.")

page_break(doc)
heading(doc, "5. Teacher salaries")
paragraph(doc, "One obligation is generated per active teacher per month. The full monthly salary is used; there is no proration.")
grid(doc, ["When", "Status", "Admin action"], [
    ("From month start through month end", "Pending", "Review amount and teacher"),
    ("After the month-end due date", "Overdue", "Open Expenses and record payment"),
    ("After payment", "Paid", "Download salary receipt"),
], [2.25, 1.25, 3.0])
step_table(doc, [
    ("Open Finance → Expenses", "Salary rows are labelled TeacherSalary and appear beside general expenses."),
    ("Choose Mark paid", "Enter payment date, method, and optional transaction/reference."),
    ("Confirm", "Status becomes Paid and the audit metadata is stored."),
    ("Download receipt", "Use the row action to save or share the teacher salary PDF."),
])

page_break(doc)
heading(doc, "6. General and recurring expenses")
paragraph(doc, "Use Expenses for rent, utilities, supplies, maintenance, marketing, and other institutional costs.")
grid(doc, ["Option", "Use it when", "Effect"], [
    ("Paid now", "Money has already left the institute", "Creates a Paid expense with payment metadata"),
    ("Pay later", "The invoice is due in future", "Creates Pending, then displays Overdue at due date"),
    ("Recurring", "The same obligation repeats", "Creates one unique occurrence per cycle"),
], [1.25, 2.65, 2.6])
heading(doc, "Recurring expense frequencies", 2)
paragraph(doc, "Monthly, Quarterly, Half-yearly, and Annual recurrence use the same idempotent occurrence rules as fees. Pending obligations count in accrual-style expenses and net profit.")
step_table(doc, [
    ("Add Expense", "Enter category, description, amount, payee, date, and due date."),
    ("Choose payment timing", "Select Paid now or Pay later."),
    ("Optional recurrence", "Choose frequency and recurrence end date."),
    ("Pay later obligation", "Use Mark paid when settled, then download the PDF receipt."),
])

page_break(doc)
heading(doc, "7. Reports, reminders, and troubleshooting")
heading(doc, "Reading the totals", 2)
grid(doc, ["Metric", "Meaning"], [
    ("Revenue", "Completed fee payments only"),
    ("Pending", "Outstanding student installment balances"),
    ("Expenses", "General expenses plus teacher salary obligations"),
    ("Net profit", "Revenue minus accrual expenses"),
], [1.45, 5.05])
heading(doc, "Troubleshooting", 2)
grid(doc, ["Symptom", "Check"], [
    ("Student is missing during fee setup", "Confirm the selected academic session and that the student is active."),
    ("A fee is not overdue yet", "Check its next-cycle due date; pending is correct before that date."),
    ("Delete fee is blocked", "Paid/partially paid records are immutable. Only unpaid rows without payments can be deleted."),
    ("Parent reminder remains", "Confirm the installment balance is zero, waived, or cancelled."),
    ("Email receipt fails", "The parent must add a valid recovery email; provider failure is reported, not hidden."),
    ("Duplicate-looking charge", "Compare period start and occurrence key. Genuine duplicates should be reported to support."),
], [2.15, 4.35])
callout(doc, "Safe operating rule", "Use production students such as test/test2 read-only during verification. Create payments only against isolated fixtures so financial reporting remains trustworthy.", fill="FFF8E8", accent=GOLD)

OUT.mkdir(parents=True, exist_ok=True)
doc.save(DOCX)
print(DOCX)
